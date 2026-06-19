#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Автоматические тесты для BookToText."""

import os
import sys
import tempfile
import shutil
import zipfile

# Добавляем путь проекта
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from converters.text_cleaner import clean_text, safe_save_path
from converters.epub_converter import convert_epub
from converters.fb2_converter import convert_fb2
from converters.pdf_converter import convert_pdf
from converters.docx_converter import convert_docx
from converters.rtf_converter import convert_rtf
from converters.html_converter import convert_html, html_to_markdown
from converters.metadata_extractor import sanitize_filename, parse_metadata_from_filename, extract_metadata


class TestResults:
    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.errors = []

    def ok(self, name):
        self.passed += 1
        print(f"  ✅ {name}")

    def fail(self, name, reason):
        self.failed += 1
        self.errors.append((name, reason))
        print(f"  ❌ {name}: {reason}")

    def summary(self):
        total = self.passed + self.failed
        print(f"\n{'='*60}")
        print(f"Результат: {self.passed}/{total} тестов пройдено")
        if self.errors:
            print(f"\nОшибки:")
            for name, reason in self.errors:
                print(f"  - {name}: {reason}")
        else:
            print("Все тесты пройдены успешно! 🎉")
        print(f"{'='*60}")
        return self.failed == 0


def test_clean_text(results: TestResults):
    """Тесты модуля очистки текста."""
    print("\n--- Тесты: clean_text ---")

    # Тест 1: Пустая строка
    out = clean_text("")
    if out == "":
        results.ok("Пустая строка")
    else:
        results.fail("Пустая строка", f"Ожидалось '', получено '{out}'")

    # Тест 2: Удаление концевых пробелов
    out = clean_text("Привет   \nМир  ")
    if "   " not in out.split("\n")[0]:
        results.ok("Удаление концевых пробелов")
    else:
        results.fail("Удаление концевых пробелов", f"Пробелы не удалены")

    # Тест 3: Дублирующиеся пустые строки (макс 2)
    out = clean_text("Строка 1\n\n\n\n\n\nСтрока 2")
    max_empty = 0
    current_empty = 0
    for line in out.splitlines():
        if line == '':
            current_empty += 1
            max_empty = max(max_empty, current_empty)
        else:
            current_empty = 0
    if max_empty <= 2:
        results.ok("Ограничение пустых строк (≤2)")
    else:
        results.fail("Ограничение пустых строк", f"Найдено {max_empty} подряд")

    # Тест 4: Удаление управляющих символов
    out = clean_text("Тест\x01\x02\x03символов")
    if '\x01' not in out and '\x02' not in out:
        results.ok("Удаление управляющих символов")
    else:
        results.fail("Удаление управляющих символов", "Символы не удалены")

    # Тест 5: Нормализация \r\n
    out = clean_text("Строка 1\r\nСтрока 2\rСтрока 3")
    lines = out.strip().split('\n')
    if len(lines) == 3 and '\r' not in out:
        results.ok("Нормализация переводов строк (\\r\\n)")
    else:
        results.fail("Нормализация переводов строк",
                     f"Строк: {len(lines)}, \\r найден: {'\\r' in out}")

    # Тест 6: Только пробелы/табы
    out = clean_text("   \t  \n  \t  ")
    if out == "":
        results.ok("Строка из одних пробелов")
    else:
        results.fail("Строка из одних пробелов", f"Получено: '{out}'")


def test_safe_save_path(results: TestResults):
    """Тесты генерации безопасного пути с умным переименованием."""
    print("\n--- Тесты: safe_save_path ---")

    with tempfile.TemporaryDirectory() as tmpdir:
        # Тест 1: Файл не существует (классика)
        path = os.path.join(tmpdir, "book.epub")
        out = safe_save_path(path)
        expected = os.path.join(tmpdir, "book.txt")
        if out == expected:
            results.ok("Путь без конфликта (классика)")
        else:
            results.fail("Путь без конфликта", f"{out} != {expected}")

        # Тест 2: Файл существует — добавляем индекс
        with open(expected, 'w') as f:
            f.write("test")
        out = safe_save_path(path)
        expected2 = os.path.join(tmpdir, "book_1.txt")
        if out == expected2:
            results.ok("Числовой индекс при конфликте")
        else:
            results.fail("Числовой индекс", f"{out} != {expected2}")

        # Тест 3: Умное переименование (Автор + Название + .md)
        out = safe_save_path(path, author="Пушкин", title="Евгений Онегин", ext=".md")
        expected_md = os.path.join(tmpdir, "Пушкин - Евгений Онегин.md")
        if out == expected_md:
            results.ok("Умное переименование (Автор - Название.md)")
        else:
            results.fail("Умное переименование", f"{out} != {expected_md}")
            
        # Тест 4: Умное переименование (только название, без точки в ext)
        out = safe_save_path(path, title="Название", ext="md")
        expected_md2 = os.path.join(tmpdir, "Название.md")
        if out == expected_md2:
            results.ok("Умное переименование (только название)")
        else:
            results.fail("Умное переименование (название)", f"{out} != {expected_md2}")

        # Тест 5: FB2.ZIP — двойное расширение
        fb2zip_path = os.path.join(tmpdir, "novel.fb2.zip")
        out = safe_save_path(fb2zip_path)
        expected_fb2 = os.path.join(tmpdir, "novel.txt")
        if out == expected_fb2:
            results.ok("FB2.ZIP двойное расширение → novel.txt")
        else:
            results.fail("FB2.ZIP двойное расширение", f"{out} != {expected_fb2}")


def test_broken_files(results: TestResults):
    """Тесты обработки сломанных/пустых файлов."""
    print("\n--- Тесты: сломанные и пустые файлы ---")

    with tempfile.TemporaryDirectory() as tmpdir:
        # Пустой EPUB
        empty_epub = os.path.join(tmpdir, "empty.epub")
        with open(empty_epub, 'wb') as f:
            f.write(b'')
        try:
            convert_epub(empty_epub)
            results.fail("Пустой EPUB", "Не вызвал исключение")
        except Exception as e:
            results.ok(f"Пустой EPUB → исключение: {type(e).__name__}")

        # Битый PDF (случайные байты)
        broken_pdf = os.path.join(tmpdir, "broken.pdf")
        with open(broken_pdf, 'wb') as f:
            f.write(b'NOT A PDF FILE CONTENT 12345')
        try:
            convert_pdf(broken_pdf)
            results.fail("Битый PDF", "Не вызвал исключение")
        except Exception as e:
            results.ok(f"Битый PDF → исключение: {type(e).__name__}")

        # Битый FB2
        broken_fb2 = os.path.join(tmpdir, "broken.fb2")
        with open(broken_fb2, 'w') as f:
            f.write("NOT XML AT ALL <>")
        try:
            convert_fb2(broken_fb2)
            results.fail("Битый FB2", "Не вызвал исключение")
        except Exception as e:
            results.ok(f"Битый FB2 → исключение: {type(e).__name__}")

        # Битый DOCX
        broken_docx = os.path.join(tmpdir, "broken.docx")
        with open(broken_docx, 'wb') as f:
            f.write(b'not a docx')
        try:
            convert_docx(broken_docx)
            results.fail("Битый DOCX", "Не вызвал исключение")
        except Exception as e:
            results.ok(f"Битый DOCX → исключение: {type(e).__name__}")

        # Пустой RTF
        empty_rtf = os.path.join(tmpdir, "empty.rtf")
        with open(empty_rtf, 'w') as f:
            f.write("")
        try:
            text = convert_rtf(empty_rtf)
            results.ok(f"Пустой RTF → не упал (текст: '{text[:30]}')")
        except Exception as e:
            results.ok(f"Пустой RTF → исключение: {type(e).__name__}")

        # Валидный HTML
        valid_html = os.path.join(tmpdir, "test.html")
        with open(valid_html, 'w', encoding='utf-8') as f:
            f.write('<html><head><style>body{color:red}</style>'
                    '<script>alert(1)</script></head>'
                    '<body><p>Привет мир</p></body></html>')
        try:
            text = convert_html(valid_html)
            if ('Привет мир' in text and 'alert' not in text
                    and 'color:red' not in text):
                results.ok("HTML: очистка от JS/CSS")
            else:
                results.fail("HTML: очистка", f"Текст: {text[:100]}")
        except Exception as e:
            results.fail("HTML: очистка", str(e))

        # Валидный RTF
        valid_rtf = os.path.join(tmpdir, "test.rtf")
        with open(valid_rtf, 'w') as f:
            f.write(r'{\rtf1\ansi Hello World}')
        try:
            text = convert_rtf(valid_rtf)
            if 'Hello World' in text:
                results.ok("RTF: извлечение текста")
            else:
                results.fail("RTF: извлечение", f"Текст: {text[:100]}")
        except Exception as e:
            results.fail("RTF: извлечение", str(e))

        # Валидный FB2
        valid_fb2 = os.path.join(tmpdir, "test.fb2")
        with open(valid_fb2, 'w', encoding='utf-8') as f:
            f.write('<?xml version="1.0" encoding="utf-8"?>'
                    '<FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0">'
                    '<body><section><p>Тестовый текст</p></section></body>'
                    '</FictionBook>')
        try:
            text = convert_fb2(valid_fb2)
            if 'Тестовый текст' in text:
                results.ok("FB2: извлечение текста")
            else:
                results.fail("FB2: извлечение", f"Текст: {text[:100]}")
        except Exception as e:
            results.fail("FB2: извлечение", str(e))

        # FB2.ZIP — архив с валидным fb2
        fb2zip_path = os.path.join(tmpdir, "test.fb2.zip")
        fb2_content = ('<?xml version="1.0" encoding="utf-8"?>'
                       '<FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0">'
                       '<body><section><p>Текст из архива</p></section></body>'
                       '</FictionBook>')
        with zipfile.ZipFile(fb2zip_path, 'w') as zf:
            zf.writestr("book.fb2", fb2_content)
        try:
            text = convert_fb2(fb2zip_path)
            if 'Текст из архива' in text:
                results.ok("FB2.ZIP: извлечение текста из архива")
            else:
                results.fail("FB2.ZIP: извлечение", f"Текст: {text[:100]}")
        except Exception as e:
            results.fail("FB2.ZIP: извлечение", str(e))

        # FB2.ZIP — пустой архив без .fb2 внутри
        empty_fb2zip = os.path.join(tmpdir, "empty.fb2.zip")
        with zipfile.ZipFile(empty_fb2zip, 'w') as zf:
            zf.writestr("readme.txt", "not a fb2")
        try:
            convert_fb2(empty_fb2zip)
            results.fail("FB2.ZIP без .fb2", "Не вызвал исключение")
        except Exception as e:
            results.ok(f"FB2.ZIP без .fb2 → исключение: {type(e).__name__}")


def test_encoding(results: TestResults):
    """Тесты UTF-8 кодировки."""
    print("\n--- Тесты: кодировка UTF-8 ---")

    # Тест с кириллицей
    text = clean_text("Привет мир! Ёжик в тумане — «кавычки»")
    if 'Привет' in text and 'Ёжик' in text:
        results.ok("Кириллица UTF-8")
    else:
        results.fail("Кириллица UTF-8", "Символы потеряны")

    # Тест с эмодзи и спецсимволами
    text = clean_text("Тест 🎉 символов €£¥ ñ ü ö")
    if '€' in text and 'ñ' in text:
        results.ok("Спецсимволы UTF-8")
    else:
        results.fail("Спецсимволы UTF-8", f"Символы потеряны: {text[:50]}")

    # Тест: строка с только байтами замены
    text = clean_text("\x00\x01\x02\x03")
    if '\x00' not in text and '\x01' not in text:
        results.ok("Очистка null и управляющих байтов")
    else:
        results.fail("Очистка null/управляющих", f"Символы остались")


def test_error_isolation(results: TestResults):
    """Тест изоляции ошибок — проверка определения форматов."""
    print("\n--- Тесты: изоляция ошибок ---")

    # Маппинг расширений → конвертеров (повторяет логику main.py)
    # Для mobi/azw используем заглушку (lambda), т.к. сам mobi_converter не нужен для теста
    _mobi_placeholder = lambda x: x
    SUPPORTED_EXTENSIONS = {
        '.epub': convert_epub,
        '.fb2': convert_fb2,
        '.pdf': convert_pdf,
        '.docx': convert_docx,
        '.mobi': _mobi_placeholder,
        '.azw3': _mobi_placeholder,
        '.azw': _mobi_placeholder,
        '.rtf': convert_rtf,
        '.html': convert_html,
        '.htm': convert_html,
    }
    FB2_ZIP = '.fb2.zip'

    def get_converter(file_path: str):
        lower = file_path.lower()
        if lower.endswith(FB2_ZIP):
            return convert_fb2
        ext = os.path.splitext(lower)[1]
        return SUPPORTED_EXTENSIONS.get(ext)

    # Проверяем, что все поддерживаемые форматы находятся
    test_cases = [
        ("book.epub", True),
        ("book.fb2", True),
        ("book.fb2.zip", True),
        ("book.pdf", True),
        ("book.docx", True),
        ("book.mobi", True),
        ("book.azw3", True),
        ("book.rtf", True),
        ("page.html", True),
        ("page.htm", True),
        ("file.xyz", False),
        ("file.exe", False),
        ("BOOK.EPUB", True),  # верхний регистр
        ("archive.FB2.ZIP", True),  # верхний регистр fb2.zip
    ]

    all_ok = True
    for filename, should_find in test_cases:
        converter = get_converter(f"/tmp/{filename}")
        found = converter is not None
        if found != should_find:
            all_ok = False
            results.fail(f"Формат {filename}",
                         f"Ожидалось {'найден' if should_find else 'не найден'}, "
                         f"получено {'найден' if found else 'не найден'}")

    if all_ok:
        results.ok("Определение всех форматов корректно (14 вариантов)")

    # Тест: обработка несуществующего файла не роняет приложение
    try:
        convert_epub("/tmp/nonexistent_file_12345.epub")
        results.fail("Несуществующий файл", "Не вызвал исключение")
    except Exception as e:
        results.ok(f"Несуществующий файл → исключение: {type(e).__name__}")


def test_metadata_extractor(results: TestResults):
    """Тесты парсинга имени файла и очистки символов."""
    print("\n--- Тесты: metadata_extractor ---")
    
    # Sanitize
    name = sanitize_filename('Имя: со слэшами/\\ и "кавычками"?!')
    if name == 'Имя со слэшами и кавычками!':
        results.ok("Очистка недопустимых символов в имени")
    else:
        results.fail("Очистка символов", f"Получено: {name}")
        
    # Разбор имени: Автор - Название
    author, title = parse_metadata_from_filename("Толстой Л.Н. - Война и мир.epub")
    if author == "Толстой Л.Н." and title == "Война и мир":
        results.ok("Разбор имени: Автор - Название")
    else:
        results.fail("Разбор имени", f"Получено: {author}, {title}")
        
    # Разбор имени: Без автора
    author, title = parse_metadata_from_filename("Просто название (ru) [litres].pdf")
    if author is None and title == "Просто название":
        results.ok("Разбор имени: очистка технического мусора")
    else:
        results.fail("Разбор имени (без автора)", f"Получено: {author}, {title}")


def test_markdown_generation(results: TestResults):
    """Тесты генерации Markdown."""
    print("\n--- Тесты: генерация Markdown ---")
    
    from bs4 import BeautifulSoup
    html = "<body><h1>Глава 1</h1><p><b>Жирный</b> и <i>курсив</i></p></body>"
    soup = BeautifulSoup(html, 'html.parser')
    md = html_to_markdown(soup.body).strip()
    
    if "# Глава 1" in md and "**Жирный**" in md and "*курсив*" in md:
        results.ok("html_to_markdown генерация разметки")
    else:
        results.fail("html_to_markdown", f"Некорректно сгенерирован MD:\n{md}")


def test_advanced_parsers(results: TestResults):
    """Тесты улучшенных конвертеров."""
    print("\n--- Тесты: Продвинутые конвертеры ---")
    
    with tempfile.TemporaryDirectory() as tmpdir:
        # Тест: PDF hyphens
        from converters.pdf_converter import clean_pdf_hyphens
        cleaned = clean_pdf_hyphens("инте-\nресный")
        if cleaned == "интересный":
            results.ok("Удаление переносов строк (PDF)")
        else:
            results.fail("Удаление переносов (PDF)", f"Получено: {cleaned}")
            
        # Тест: FB2 poem
        from converters.fb2_converter import _parse_fb2_content
        fb2_xml = """<?xml version="1.0" encoding="utf-8"?>
        <FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0">
            <body><poem><stanza><v>Стих 1</v><v>Стих 2</v></stanza></poem></body>
        </FictionBook>
        """.encode('utf-8')
        md_text = _parse_fb2_content(fb2_xml, format_mode='md')
        if "> Стих 1" in md_text and "> Стих 2" in md_text:
            results.ok("FB2: Парсинг poem и stanza")
        else:
            results.fail("FB2: Парсинг poem", f"Получено: {md_text}")
            
        # Тест: DOCX table and list
        from docx import Document
        docx_path = os.path.join(tmpdir, "test.docx")
        doc = Document()
        doc.add_paragraph("List item", style="List Bullet")
        table = doc.add_table(rows=1, cols=2)
        cells = table.rows[0].cells
        cells[0].text = "Cell 1"
        cells[1].text = "Cell 2"
        doc.save(docx_path)
        
        from converters.docx_converter import convert_docx
        md_text = convert_docx(docx_path, format_mode='md')
        if "- List item" in md_text and "| Cell 1 | Cell 2 |" in md_text:
            results.ok("DOCX: Парсинг списков и таблиц")
        else:
            results.fail("DOCX: Парсинг", f"Получено: {md_text}")
            
        # Тест: Кодировка CP1251
        from converters.html_converter import convert_html
        html_path = os.path.join(tmpdir, "test_cp1251.html")
        with open(html_path, 'wb') as f:
            f.write("<html><body><p>Привет, мир! Это довольно длинное сообщение на русском языке для определения кодировки.</p></body></html>".encode('cp1251'))
        text = convert_html(html_path)
        if "Привет" in text and "сообщение" in text:
            results.ok("HTML: Автоопределение кодировки CP1251")
        else:
            results.fail("HTML: Автоопределение кодировки", f"Получено: {text}")


def main():
    print("=" * 60)
    print("  BookToText — Автоматические тесты")
    print("=" * 60)

    results = TestResults()

    test_clean_text(results)
    test_safe_save_path(results)
    test_broken_files(results)
    test_encoding(results)
    test_error_isolation(results)
    test_metadata_extractor(results)
    test_markdown_generation(results)
    test_advanced_parsers(results)

    success = results.summary()
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
